# dashboard.py
import streamlit as st
from docx import Document
from io import BytesIO
import re

section_titles = [
    "Project Summary",
    "Challenges and Needs",
    "Target Group",
    "Organisation Structure",
    "Risk Analysis",
    "Communication Plan",
    "Internal Policies",
]


def clean_content(content):
    """Remove 'Your input' and 'AI-generated draft' labels from content"""
    if not content:
        return ""

    # Remove "Your input" sections
    content = re.sub(r"\*\*Your input:\*\*.*?\n\n", "", content, flags=re.DOTALL)

    # Remove "AI-generated draft" labels
    content = re.sub(r"\*\*AI-generated draft for .*?:\*\*\n\n", "", content)

    return content.strip()


def format_section_content(content):
    """Helper function to format section content with proper markdown"""
    cleaned = clean_content(content)
    if not cleaned:
        return "*No content provided for this section*"

    # Check if content contains markdown tables and preserve them
    if "|" in cleaned and "-" in cleaned:
        return cleaned
    return cleaned


def dashboard_ui():
    user = st.session_state.get("user", "guest@example.com")

    # Header
    st.markdown(
        f"""
        <div style="background-color: #f5f5f5; padding: 1rem; border-radius: 8px; margin-bottom: 1.5rem;">
            <h2>📋 ERDF Application Dashboard</h2>
            <p>👤 Logged in as: <strong>{user}</strong></p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Sidebar Navigation
    st.sidebar.title("Navigation")
    selected_section = st.sidebar.radio(
        "Sections", section_titles + ["Full Document Preview"]
    )

    if selected_section == "Full Document Preview":
        st.subheader("📄 Complete Application Document")
        st.markdown("---")

        # Combined document view
        with st.container():
            # Project Summary
            st.markdown("## 1. Project Summary")
            summary = st.session_state.edited_sections.get(
                "Project Summary", st.session_state.get("step_0_generated", "")
            )
            st.markdown(format_section_content(summary))
            st.markdown("---")

            # Challenges and Needs
            st.markdown("## 2. Challenges and Needs")
            challenges = st.session_state.edited_sections.get(
                "Challenges and Needs", st.session_state.get("step_0_generated", "")
            )
            st.markdown(format_section_content(challenges))
            st.markdown("---")

            # Target Group
            st.markdown("## 3. Target Group")
            target_group = st.session_state.edited_sections.get(
                "Target Group", st.session_state.get("step_2_generated", "")
            )
            st.markdown(format_section_content(target_group))
            st.markdown("---")

            # Organisation Structure
            st.markdown("## 4. Organisation Structure")
            org_structure = st.session_state.edited_sections.get(
                "Organisation Structure", st.session_state.get("step_4_generated", "")
            )
            st.markdown(format_section_content(org_structure))
            st.markdown("---")

            # Risk Analysis
            st.markdown("## 5. Risk Analysis")
            risks = st.session_state.edited_sections.get(
                "Risk Analysis", st.session_state.get("step_5_generated", "")
            )
            st.markdown(format_section_content(risks))
            st.markdown("---")

            # Communication Plan
            st.markdown("## 6. Communication Plan")
            comm_plan = st.session_state.edited_sections.get(
                "Communication Plan", st.session_state.get("step_6_generated", "")
            )
            st.markdown(format_section_content(comm_plan))
            st.markdown("---")

            # Internal Policies
            st.markdown("## 7. Internal Policies")
            policies = st.session_state.edited_sections.get(
                "Internal Policies", st.session_state.get("step_7_generated", "")
            )
            st.markdown(format_section_content(policies))

        # Edit controls
        st.markdown("---")
        with st.expander("✏️ Edit Sections"):
            selected_edit = st.selectbox("Select section to edit:", section_titles)
            section_index = section_titles.index(selected_edit)

            content = st.session_state.edited_sections.get(
                selected_edit,
                st.session_state.get(f"step_{section_index}_generated", ""),
            )

            edited = st.text_area(
                f"Edit {selected_edit}",
                value=content,
                height=200,
                key=f"full_preview_edit_{section_index}",
            )

            if st.button(f"💾 Save changes to {selected_edit}"):
                st.session_state.edited_sections[selected_edit] = edited
                st.success(f"Changes to {selected_edit} saved!")
                st.rerun()

        # Export options
        st.markdown("---")
        st.subheader("Export Options")

        if st.button("⬇️ Download as DOCX"):
            doc = Document()
            doc.add_heading("ERDF Application", 0)

            sections = [
                ("1. Project Summary", "Project Summary", 0),
                ("2. Challenges and Needs", "Challenges and Needs", 0),
                ("3. Target Group", "Target Group", 2),
                ("4. Organisation Structure", "Organisation Structure", 4),
                ("5. Risk Analysis", "Risk Analysis", 5),
                ("6. Communication Plan", "Communication Plan", 6),
                ("7. Internal Policies", "Internal Policies", 7),
            ]

            for title, section_name, step in sections:
                doc.add_heading(title, level=1)
                content = st.session_state.edited_sections.get(
                    section_name, st.session_state.get(f"step_{step}_generated", "")
                )
                cleaned_content = clean_content(content)

                # Handle tables in the Word document
                if "|" in cleaned_content and "-" in cleaned_content:
                    # This is a simple table - for better table handling you might need more complex parsing
                    table = doc.add_table(rows=1, cols=1)
                    table.cell(0, 0).text = cleaned_content
                else:
                    doc.add_paragraph(cleaned_content)

            buffer = BytesIO()
            doc.save(buffer)
            st.download_button(
                "Download DOCX",
                buffer.getvalue(),
                file_name="ERDF_Application.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        return

    # Individual section view
    st.subheader(f"✏️ {selected_section}")

    # Get the section index
    section_index = section_titles.index(selected_section)

    # Get the original or edited content
    content = st.session_state.edited_sections.get(
        selected_section,
        st.session_state.get(
            f"step_{section_index}_generated", "No content available yet."
        ),
    )

    # Large editable text area
    edited_content = st.text_area(
        "Edit this section:", value=content, height=400, key=f"edit_{selected_section}"
    )

    # Save button
    if st.button("💾 Save Changes"):
        st.session_state.edited_sections[selected_section] = edited_content
        st.success("Changes saved to your application!")

    # Show related user input (only in edit mode, not in final doc)
    st.divider()
    st.subheader("Your Original Input")
    user_input = st.session_state.get(
        f"step_{section_index}_input", "No input provided."
    )
    st.info(user_input)
